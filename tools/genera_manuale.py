"""
snap - Generazione dei manuali software in formato .docx.

Converte un documento Markdown della cartella docs/ in un file Word conforme
alla convenzione tipografica di progetto: carattere PT Sans Narrow a 19 punti e
stili predefiniti di Word (Titolo 1..3, Corpo del testo, Griglia tabella).

Uso:
    python tools/genera_manuale.py                     tutti i documenti di docs/
    python tools/genera_manuale.py 05_MANUALE_OPERATIVO.md
    python tools/genera_manuale.py --uscita C:\\percorso

remarks: Autore: Daniele Speziale - Data: 2026-08-27
copyright: (c) 2024-26 DS Consulting
license: MIT
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

RADICE = Path(__file__).resolve().parent.parent
DOCS = RADICE / "docs"

CARATTERE = "PT Sans Narrow"
CORPO_PT = 19
# I titoli scalano a partire dal corpo, restando nella stessa famiglia.
DIMENSIONI_TITOLI = {1: 26, 2: 23, 3: 21, 4: 19, 5: 19, 6: 19}


def _imposta_carattere(esecuzione, dimensione: int, corsivo: bool = False,
                       grassetto: bool = False, monospazio: bool = False) -> None:
    esecuzione.font.name = "Consolas" if monospazio else CARATTERE
    esecuzione.font.size = Pt(dimensione if not monospazio else dimensione - 4)
    esecuzione.font.italic = corsivo
    esecuzione.font.bold = grassetto
    # Word richiede anche la dichiarazione per gli alfabeti non latini.
    proprieta = esecuzione._element.get_or_add_rPr()
    caratteri = proprieta.get_or_add_rFonts()
    caratteri.set(qn("w:eastAsia"), CARATTERE)


def _applica_stili_documento(documento: Document) -> None:
    """Porta gli stili predefiniti di Word alla convenzione di progetto."""
    for nome, dimensione in [
        ("Normal", CORPO_PT),
        ("Body Text", CORPO_PT),
        ("List Bullet", CORPO_PT),
        ("List Number", CORPO_PT),
        ("Caption", CORPO_PT - 4),
        ("Heading 1", DIMENSIONI_TITOLI[1]),
        ("Heading 2", DIMENSIONI_TITOLI[2]),
        ("Heading 3", DIMENSIONI_TITOLI[3]),
        ("Heading 4", DIMENSIONI_TITOLI[4]),
        ("Title", 32),
    ]:
        try:
            stile = documento.styles[nome]
        except KeyError:
            continue  # stile non presente nel modello: si prosegue
        stile.font.name = CARATTERE
        stile.font.size = Pt(dimensione)
        proprieta = stile.element.get_or_add_rPr()
        proprieta.get_or_add_rFonts().set(qn("w:eastAsia"), CARATTERE)


def _testo_formattato(paragrafo, testo: str, dimensione: int) -> None:
    """Rende grassetto, corsivo e monospazio del Markdown in esecuzioni Word."""
    pezzi = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", testo)
    for pezzo in pezzi:
        if not pezzo:
            continue
        if pezzo.startswith("**") and pezzo.endswith("**"):
            _imposta_carattere(paragrafo.add_run(pezzo[2:-2]), dimensione, grassetto=True)
        elif pezzo.startswith("*") and pezzo.endswith("*"):
            _imposta_carattere(paragrafo.add_run(pezzo[1:-1]), dimensione, corsivo=True)
        elif pezzo.startswith("`") and pezzo.endswith("`"):
            _imposta_carattere(paragrafo.add_run(pezzo[1:-1]), dimensione, monospazio=True)
        else:
            _imposta_carattere(paragrafo.add_run(pezzo), dimensione)


def _righe_tabella(righe: list[str]) -> list[list[str]]:
    tabella = []
    for riga in righe:
        celle = [c.strip() for c in riga.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{2,}:?", c or "-") for c in celle):
            continue  # riga di separazione dell'intestazione
        tabella.append(celle)
    return tabella


def converti(sorgente: Path, destinazione: Path) -> Path:
    """Converte un documento Markdown in .docx applicando la convenzione."""
    documento = Document()
    _applica_stili_documento(documento)

    testo = sorgente.read_text(encoding="utf-8")
    # I commenti di intestazione del Markdown non fanno parte del documento.
    testo = re.sub(r"<!--.*?-->", "", testo, flags=re.DOTALL)

    righe = testo.splitlines()
    indice = 0
    in_codice = False
    blocco_codice: list[str] = []

    while indice < len(righe):
        riga = righe[indice]

        if riga.strip().startswith("```"):
            if in_codice:
                paragrafo = documento.add_paragraph()
                _imposta_carattere(paragrafo.add_run("\n".join(blocco_codice)),
                                   CORPO_PT, monospazio=True)
                paragrafo.paragraph_format.left_indent = Pt(18)
                blocco_codice = []
            in_codice = not in_codice
            indice += 1
            continue

        if in_codice:
            blocco_codice.append(riga)
            indice += 1
            continue

        # Tabelle
        if riga.strip().startswith("|"):
            gruppo = []
            while indice < len(righe) and righe[indice].strip().startswith("|"):
                gruppo.append(righe[indice])
                indice += 1
            dati = _righe_tabella(gruppo)
            if dati:
                tabella = documento.add_table(rows=len(dati), cols=len(dati[0]))
                tabella.style = "Table Grid"
                for numero, contenuto in enumerate(dati):
                    for colonna, valore in enumerate(contenuto[: len(dati[0])]):
                        cella = tabella.cell(numero, colonna)
                        cella.text = ""
                        paragrafo = cella.paragraphs[0]
                        _testo_formattato(paragrafo, valore, CORPO_PT - 4)
                        if numero == 0:
                            for esecuzione in paragrafo.runs:
                                esecuzione.font.bold = True
                documento.add_paragraph()
            continue

        # Titoli
        titolo = re.match(r"^(#{1,6})\s+(.*)$", riga)
        if titolo:
            livello = len(titolo.group(1))
            paragrafo = documento.add_heading(level=min(livello, 4))
            paragrafo.text = ""
            _testo_formattato(paragrafo, titolo.group(2), DIMENSIONI_TITOLI.get(livello, CORPO_PT))
            for esecuzione in paragrafo.runs:
                esecuzione.font.bold = True
            indice += 1
            continue

        # Separatori
        if re.fullmatch(r"-{3,}", riga.strip()):
            separatore = documento.add_paragraph()
            separatore.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _imposta_carattere(separatore.add_run("* * *"), CORPO_PT)
            indice += 1
            continue

        # Elenchi
        elenco = re.match(r"^\s*[-*]\s+(.*)$", riga)
        numerato = re.match(r"^\s*\d+\.\s+(.*)$", riga)
        if elenco or numerato:
            stile = "List Bullet" if elenco else "List Number"
            paragrafo = documento.add_paragraph(style=stile)
            _testo_formattato(paragrafo, (elenco or numerato).group(1), CORPO_PT)
            indice += 1
            continue

        if not riga.strip():
            indice += 1
            continue

        # Paragrafo ordinario: si accorpano le righe consecutive.
        blocco = [riga.strip()]
        indice += 1
        while (indice < len(righe) and righe[indice].strip()
               and not re.match(r"^(#{1,6}\s|\||```|\s*[-*]\s|\s*\d+\.\s|-{3,})", righe[indice])):
            blocco.append(righe[indice].strip())
            indice += 1
        paragrafo = documento.add_paragraph()
        _testo_formattato(paragrafo, " ".join(blocco), CORPO_PT)

    destinazione.parent.mkdir(parents=True, exist_ok=True)
    documento.save(str(destinazione))
    return destinazione


def main() -> int:
    parser = argparse.ArgumentParser(description="Genera i manuali software in .docx")
    parser.add_argument("documenti", nargs="*", help="nomi dei file in docs/ (predefinito: tutti)")
    parser.add_argument("--uscita", default=str(DOCS / "docx"), help="cartella di destinazione")
    argomenti = parser.parse_args()

    if argomenti.documenti:
        sorgenti = [DOCS / nome for nome in argomenti.documenti]
    else:
        sorgenti = sorted(DOCS.glob("*.md"))

    mancanti = [s for s in sorgenti if not s.exists()]
    if mancanti:
        for s in mancanti:
            print("documento non trovato: %s" % s, file=sys.stderr)
        return 2

    uscita = Path(argomenti.uscita)
    print("Carattere: %s %dpt (stili predefiniti di Word)" % (CARATTERE, CORPO_PT))
    for sorgente in sorgenti:
        prodotto = converti(sorgente, uscita / (sorgente.stem + ".docx"))
        print("  %-34s -> %s (%d kB)"
              % (sorgente.name, prodotto.name, prodotto.stat().st_size // 1024))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
